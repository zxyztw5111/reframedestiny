#!/usr/bin/env python3
"""Export course notes markdown to Word (.docx) for WPS Office."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import sys

MD_PATH = "/Users/zhangxinyun/Projects/reframe-destiny/research/course-notes-week1-7.md"
OUT_PATH = "/Users/zhangxinyun/Desktop/AI北大/Generation-AI-Week1-7-课程笔记-张昕田.docx"


def set_chinese_font(run, name="宋体", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    if level == 1:
        set_chinese_font(run, "黑体", 18)
        run.bold = True
    elif level == 2:
        set_chinese_font(run, "黑体", 14)
        run.bold = True
    else:
        set_chinese_font(run, "黑体", 12)
        run.bold = True
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, size=size)
    run.bold = bold
    run.italic = italic
    return p


def parse_table_lines(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-:\s|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < ncols:
                cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
                cell_text = re.sub(r"`(.+?)`", r"\1", cell_text)
                cell_text = cell_text.replace("✅", "✓").replace("❌", "✗").replace("🟡", "△").replace("⬜", "□")
                table.rows[i].cells[j].text = cell_text
                for para in table.rows[i].cells[j].paragraphs:
                    for run in para.runs:
                        set_chinese_font(run, size=10)
    doc.add_paragraph()


def convert_md_to_docx(md_path, out_path):
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split("\n")
    i = 0
    in_table = False
    table_buf = []

    while i < len(lines):
        line = lines[i]

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = add_paragraph(doc, text, italic=True, size=10)
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,4})\s+(.+)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            title = m.group(2).strip()
            title = re.sub(r"\*\*(.+?)\*\*", r"\1", title)
            add_heading(doc, title, level)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Table
        if line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_buf = []
            table_buf.append(line)
            i += 1
            # Check if next line is still table
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                rows = parse_table_lines(table_buf)
                add_table(doc, rows)
                in_table = False
                table_buf = []
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Italic-only line
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            text = line.strip().strip("*")
            add_paragraph(doc, text, italic=True, size=10)
            i += 1
            continue

        # Normal paragraph
        text = line.strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        add_paragraph(doc, text)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    md = sys.argv[1] if len(sys.argv) > 1 else MD_PATH
    out = sys.argv[2] if len(sys.argv) > 2 else OUT_PATH
    convert_md_to_docx(md, out)
